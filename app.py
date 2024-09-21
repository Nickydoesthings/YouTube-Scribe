from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from flask_wtf import FlaskForm, CSRFProtect
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired, Email, EqualTo, Length
from flask_migrate import Migrate
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_caching import Cache
import os
import yt_dlp
from openai import OpenAI
import logging
import markdown
from io import BytesIO, StringIO
from weasyprint import HTML
import webvtt
from datetime import timedelta, datetime
from docx import Document
from bs4 import BeautifulSoup  # To parse HTML
import sib_api_v3_sdk  # For Brevo stuff
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature


# Initialize the Flask application
app = Flask(__name__)

cache_config = {
    "CACHE_TYPE": "filesystem",  # Use file-based cache
    "CACHE_DIR": "cache-directory",  # Specify the directory where the cache will be stored
}

cache = Cache(app, config=cache_config)

# Initialize Flask-Limiter for rate limiting
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"],  # Applying the default limits globally
    storage_uri="memory://",  # Uses memory cache, but data will persist in files
    storage_options={"cache": cache}  # Use Flask-Caching for storage
)
# Configuration for database and login management
app.config['SECRET_KEY'] = 'your_secret_key_here'
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize extensions
db = SQLAlchemy(app)
migrate = Migrate(app, db)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Initialize CSRF protection
csrf = CSRFProtect(app)

# Set up logging to track and debug any issues during runtime
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the OpenAI client using the API key from environment variables
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Initialize the serializer for tokens
serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])


# Define User model
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(150), nullable=False)
    is_confirmed = db.Column(db.Boolean, default=False)
    confirmation_token = db.Column(db.String(100), nullable=True)
    usage_count = db.Column(db.Integer, default=0)
    plan = db.Column(db.String(50), default='free')
    last_confirmation_sent_at = db.Column(db.DateTime, nullable=True)
    failed_attempts = db.Column(db.Integer, default=0)  # Track failed login attempts
    lock_until = db.Column(db.DateTime, nullable=True)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Define the forms
class SignupForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email(), Length(max=150)])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password', message="Passwords must match.")])
    submit = SubmitField('Sign Up')

class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email(), Length(max=150)])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Log In')

class ResendConfirmationForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email(), Length(max=150)])
    submit = SubmitField('Resend Confirmation Email')

class ResetPasswordModalForm(FlaskForm):
    email = StringField('Email', render_kw={'readonly': True})
    submit = SubmitField('Send Reset Link')

class ResetPasswordRequestForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email(), Length(max=150)])
    submit = SubmitField('Send Reset Link')

class ResetPasswordForm(FlaskForm):
    password = PasswordField('New Password', validators=[DataRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm New Password', validators=[DataRequired(), EqualTo('password', message="Passwords must match.")])
    submit = SubmitField('Update Password')

class LogoutForm(FlaskForm):
    submit = SubmitField('Sign Out')

# for email confirmation resends
RESEND_COOLDOWN = timedelta(minutes=2)  # Cooldown for resending confirmation email after 1st resend
LOCKOUT_DURATION = timedelta(minutes=1)  # Lock user for 1 minute
MAX_ATTEMPTS = 5  # Max failed login attempts allowed

def can_resend_confirmation(user):
    if not user.last_confirmation_sent_at:
        return True
    return datetime.utcnow() - user.last_confirmation_sent_at > RESEND_COOLDOWN

# Set up global rate limit that applies to all requests
@app.before_request
@limiter.limit("200 per day; 50 per hour")
def apply_global_limit():
    pass

@login_manager.unauthorized_handler
def unauthorized():
    return redirect(url_for('login'))

# Routes for user authentication
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    form = SignupForm()
    if form.validate_on_submit():
        email = form.email.data.strip()
        password = form.password.data

        # Check if the user already exists
        user = User.query.filter_by(email=email).first()
        if user:
            flash('Email already exists.', 'danger')
            return redirect(url_for('signup'))

        # Create a new user with confirmation token
        try:
            new_user = User(email=email, password=generate_password_hash(password, method='pbkdf2:sha256'))
            db.session.add(new_user)
            db.session.commit()
            
            # Generate token and send confirmation email
            token = generate_confirmation_token(new_user.email)
            new_user.confirmation_token = token
            db.session.commit()

            send_confirmation_email(new_user.email, token)

            # Store the user's email in session to pre-fill the resend form
            session['email_for_confirmation'] = new_user.email

            # Remove the flash message to prevent duplication on email_sent.html
            return redirect(url_for('email_sent'))
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error creating user: {e}")
            flash('An error occurred while creating the account. Please try again.', 'danger')

    elif form.errors:
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"{error}", 'danger')

    return render_template('signup.html', form=form)

# Route for email_sent.html which allows users to resend confirmation email if needed
@app.route('/email_sent', methods=['GET', 'POST'])
def email_sent():
    form = ResendConfirmationForm()
    
    # Retrieve the email from the session without removing it
    email = session.get('email_for_confirmation', None)
    
    if not email:
        flash('No email found for confirmation resending. Please sign up again.', 'danger')
        return redirect(url_for('signup'))

    # Logging to debug form validation
    if request.method == 'POST':
        logger.info("POST request received.")
        
        if form.validate_on_submit():
            logger.info("Form validation passed.")
            
            user = User.query.filter_by(email=email).first()
            
            if not user:
                flash('No account found with that email.', 'danger')
                return redirect(url_for('signup'))
            
            if user.is_confirmed:
                flash('Your email is already confirmed. Please log in.', 'success')
                return redirect(url_for('login'))
            
            if not can_resend_confirmation(user):
                flash('You can resend the confirmation email again later. Please try again after some time.', 'warning')
                return redirect(url_for('email_sent'))

            try:
                # Generate a new confirmation token
                token = generate_confirmation_token(user.email)
                user.confirmation_token = token
                user.last_confirmation_sent_at = datetime.utcnow()
                db.session.commit()

                # Send the confirmation email again
                send_confirmation_email(user.email, token)

                flash('A new confirmation email has been sent. Please check your email.', 'success')
                return redirect(url_for('email_sent'))
            except Exception as e:
                flash('An error occurred while resending the confirmation email. Please try again.', 'danger')
                return redirect(url_for('email_sent'))

        # Log form errors if validation failed
        if not form.validate_on_submit():
            logger.error(f"Form validation failed. Errors: {form.errors}")
            flash('Unexpected Error. Please try again.', 'danger')

    return render_template('email_sent.html', form=form, email=email)



@app.route('/reset_password', methods=['GET', 'POST'])
def reset_password_request():
    form = ResetPasswordRequestForm()
    if form.validate_on_submit():
        email = form.email.data.strip()
        user = User.query.filter_by(email=email).first()
        
        if user:
            # Generate a secure token
            token = serializer.dumps(user.email, salt='password-reset-salt')
            send_password_reset_email(user.email, token)
            flash('A password reset link has been sent to your email.', 'info')
            logger.info(f"Password reset requested for email: {email}")
        else:
            # Inform the user that the email does not exist
            flash('The provided email address is not registered in our system.', 'danger')
            logger.warning(f"Password reset requested for non-existent email: {email}")
        
        return redirect(url_for('login'))
    
    return render_template('reset_password_request.html', form=form)

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        email = serializer.loads(token, salt='password-reset-salt', max_age=3600)  # Token valid for 1 hour
    except SignatureExpired:
        flash('The password reset link has expired.', 'danger')
        return redirect(url_for('reset_password_request'))
    except BadSignature:
        flash('The password reset link is invalid.', 'danger')
        return redirect(url_for('reset_password_request'))
    
    user = User.query.filter_by(email=email).first_or_404()
    form = ResetPasswordForm()
    
    if form.validate_on_submit():
        password = form.password.data
        confirm_password = form.confirm_password.data
        
        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return redirect(url_for('reset_password', token=token))
        
        if len(password) < 6:
            flash('Password must be at least 6 characters long.', 'warning')
            return redirect(url_for('reset_password', token=token))
        
        # Update the user's password
        user.password = generate_password_hash(password, method='pbkdf2:sha256')
        db.session.commit()
        flash('Your password has been updated. Please log in.', 'success')
        logger.info(f"Password updated for user: {email}")
        return redirect(url_for('login'))
    
    return render_template('reset_password.html', form=form, token=token)


@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        email = form.email.data
        password = form.password.data
        user = User.query.filter_by(email=email).first()

        if user:
            # Check if user is locked out
            if user.lock_until and user.lock_until > datetime.utcnow():
                flash('Too many failed attempts. Try again later.', 'danger')
                return redirect(url_for('login'))

            # Check password
            if check_password_hash(user.password, password):
                if not user.is_confirmed:
                    flash('Please confirm your email before logging in.', 'danger')
                    return redirect(url_for('login'))
                
                # Reset failed attempts and lock_until on successful login
                user.failed_attempts = 0
                user.lock_until = None
                db.session.commit()

                login_user(user)
                return redirect(url_for('generator'))
            else:
                # Increment failed attempts
                user.failed_attempts += 1

                # Lock account if max attempts exceeded
                if user.failed_attempts >= MAX_ATTEMPTS:
                    user.lock_until = datetime.utcnow() + LOCKOUT_DURATION
                    flash(f'Too many failed attempts. Account locked for {LOCKOUT_DURATION.total_seconds() // 60} minutes.', 'danger')
                else:
                    flash('Invalid email or password.', 'danger')

                db.session.commit()
        else:
            flash('Invalid email or password.', 'danger')

    return render_template('login.html', form=form)

@app.route('/confirm/<token>')
def confirm_email(token):
    try:
        email = confirm_token(token)
    except (SignatureExpired, BadSignature):
        flash('The confirmation link is invalid or has expired.', 'danger')
        return redirect(url_for('login'))

    user = User.query.filter_by(email=email).first_or_404()

    if user.is_confirmed:
        flash('Account already confirmed. Please log in.', 'success')
    else:
        user.is_confirmed = True
        user.confirmation_token = None
        user.last_confirmation_sent_at = None  # Reset the timestamp
        db.session.commit()
        flash('You have confirmed your account. Thanks!', 'success')

    return redirect(url_for('login'))

@app.route('/logout', methods=['GET', 'POST'])
@login_required
def logout():
    logout_user()
    session.pop('email_for_confirmation', None)
    flash('You have been logged out.', 'success')
    
    return redirect(url_for('login'))

@app.route('/', methods=['GET', 'POST'])
def landing():
    user_plan = None
    if current_user.is_authenticated:
        user_plan = current_user.plan
    
    if request.method == 'POST':
        youtube_link = request.form.get('youtubeLink')
        return redirect(url_for('generator') + f"?youtubeLink={youtube_link}")
    
    return render_template('landing.html', user_plan=user_plan)

@app.route('/generator', methods=['GET', 'POST'])
@login_required
def generator():
    error = None
    html_summary = None
    youtube_link = request.args.get('youtubeLink', '')
    video_title = None
    thumbnail_url = None
    show_upgrade_popup = False
    upgrade_reason = request.args.get('upgrade_reason', None)
    video_length_formatted = None  # Variable to store formatted video length

    if request.method == 'POST':
        youtube_link = request.form.get('youtubeLink')
        if not youtube_link:
            error = "Please provide a YouTube video URL."
        else:
            _, video_title, thumbnail_url, duration = download_youtube_audio(youtube_link, metadata_only=True)

            if not video_title or not thumbnail_url:
                error = "Failed to fetch video metadata."
                return render_template('generator.html', error=error, youtube_link=youtube_link)
    
            # Convert duration from seconds to "minutes:seconds" format
            video_length_formatted = str(timedelta(seconds=duration)) if duration else None

            if current_user.is_authenticated and current_user.plan == 'pro':
                max_duration = 3600  # 2 hours in seconds
            else:
                max_duration = 900  # 15 minutes in seconds

            if duration > max_duration:
                show_upgrade_popup = True
                upgrade_reason = 'video_duration'
                return render_template('generator.html', error=error, youtube_link=youtube_link, video_title=video_title, thumbnail_url=thumbnail_url, show_upgrade_popup=show_upgrade_popup, upgrade_reason=upgrade_reason, video_length=video_length_formatted)

            captions = download_youtube_captions(youtube_link)
            if captions:
                transcript = captions
            else:
                if not current_user.is_authenticated or current_user.plan == 'free':
                    show_upgrade_popup = True
                    upgrade_reason = 'no_captions'
                    return render_template('generator.html', error=error, youtube_link=youtube_link, video_title=video_title, thumbnail_url=thumbnail_url, show_upgrade_popup=show_upgrade_popup, upgrade_reason=upgrade_reason)

                audio_file_path, _, _, _ = download_youtube_audio(youtube_link)
                transcript = transcribe_audio_with_whisper_api(audio_file_path)
                if os.path.exists(audio_file_path):
                    os.remove(audio_file_path)
                    logger.info(f"Deleted audio file: {audio_file_path}")

            try:
                if transcript:
                    summary = summarize_text(transcript)
                    if summary:
                        html_summary = markdown.markdown(summary)
                        if current_user.is_authenticated:
                            current_user.usage_count += 1
                            db.session.commit()
                    else:
                        error = "Failed to generate summary."
                else:
                    error = "Failed to obtain transcript."
            except Exception as e:
                logger.error(f"An unexpected error occurred: {e}")
                error = f"An unexpected error occurred: {e}"

    return render_template('generator.html', summary=html_summary, error=error, youtube_link=youtube_link, video_title=video_title, thumbnail_url=thumbnail_url, show_upgrade_popup=show_upgrade_popup, upgrade_reason=upgrade_reason, video_length=video_length_formatted)

@app.route('/download/pdf', methods=['POST'])
@login_required
def download_pdf():
    if current_user.plan != 'pro':
        return redirect(url_for('generator', show_upgrade_popup=True, upgrade_reason='download_pdf'))

    summary = request.form.get('summary')

    if not summary:
        flash("No summary available for download. Please generate a summary first.", "danger")
        return redirect(url_for('generator'))

    html_content = f"<html><body>{markdown.markdown(summary)}</body></html>"
    pdf_file = BytesIO()
    HTML(string=html_content).write_pdf(pdf_file)
    pdf_file.seek(0)

    return send_file(pdf_file, as_attachment=True, download_name="summary.pdf")

@app.route('/download/word', methods=['POST'])
@login_required
def download_word():
    if current_user.plan != 'pro':
        return redirect(url_for('generator', show_upgrade_popup=True, upgrade_reason='download_word'))

    summary = request.form.get('summary')

    if not summary:
        flash("No summary available for download. Please generate a summary first.", "danger")
        return redirect(url_for('generator'))

    # Create a new Word document
    doc = Document()
    doc.add_heading('Summary', 0)

    # Parse the HTML content and convert it to Word-friendly format
    soup = BeautifulSoup(summary, 'html.parser')

    for element in soup.find_all():
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            # Keep bullets as bullets
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            # Turns numbered lists into bullets bc there's no way to reset the counter between sections :( whatevr
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')  

    # Save the document in memory
    word_file = BytesIO()
    doc.save(word_file)
    word_file.seek(0)

    return send_file(word_file, as_attachment=True, download_name="summary.docx")

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/pricing')
def pricing():
    return render_template('pricing.html')

@app.route('/my_account')
@login_required
def my_account():
    reset_form = ResetPasswordModalForm()
    logout_form = LogoutForm()
    return render_template(
        'my_account.html',
        email=current_user.email,
        plan=current_user.plan,
        reset_form=reset_form,
        logout_form=logout_form
    )

@app.route('/upgrade_to_pro', methods=['GET', 'POST'])
@login_required
def upgrade_to_pro():
    current_user.plan = 'pro'
    db.session.commit()
    flash('You have been upgraded to the Pro plan.', 'success')
    return redirect(url_for('pricing'))

@app.route('/fetch_metadata', methods=['POST'])
def fetch_metadata():
    youtube_link = request.json.get('youtubeLink')

    if not youtube_link:
        return jsonify({'error': 'No YouTube link provided.'}), 400

    try:
        _, video_title, thumbnail_url, duration = download_youtube_audio(youtube_link, metadata_only=True)

        if not video_title or not thumbnail_url:
            return jsonify({'error': 'Failed to fetch video metadata.'}), 500

        return jsonify({
            'video_title': video_title,
            'thumbnail_url': thumbnail_url,
            'duration': duration
        }), 200

    except Exception as e:
        logger.error(f"An error occurred while fetching metadata: {e}")
        return jsonify({'error': 'An unexpected error occurred.'}), 500

# Error handler for rate limit exceeded (429 Too Many Requests)
@app.errorhandler(429)
def ratelimit_handler(e):
    return render_template('429.html'), 429

def download_youtube_audio(video_url, save_path='.', metadata_only=False):
    try:
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': f'{save_path}/%(title)s.%(ext)s',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'postprocessor_args': [
                '-ar', '44100'
            ],
            'prefer_ffmpeg': True,
            'keepvideo': False,
            'skip_download': metadata_only,
        }

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info_dict = ydl.extract_info(video_url, download=not metadata_only)
            if not metadata_only:
                filename = ydl.prepare_filename(info_dict)
                audio_file_path = os.path.splitext(filename)[0] + ".mp3"
            else:
                audio_file_path = None
            video_title = info_dict.get('title', 'Unknown Title')
            thumbnail_url = info_dict.get('thumbnail', '')
            duration = info_dict.get('duration', 0)

        return audio_file_path, video_title, thumbnail_url, duration
    except Exception as e:
        logger.error(f"An error occurred while downloading audio: {e}")
        return None, None, None, None

def download_youtube_captions(video_url):
    try:
        ydl_opts = {
            'skip_download': True,
            'writesubtitles': True,
            'writeautomaticsub': True,
            'subtitlesformat': 'vtt',
            'subtitleslangs': ['en'],
            'outtmpl': '%(id)s.%(ext)s',
        }

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info_dict = ydl.extract_info(video_url, download=False)

            subtitles = info_dict.get('subtitles', {})
            automatic_captions = info_dict.get('automatic_captions', {})

            captions_available = False

            if 'en' in subtitles:
                captions_available = True
            elif 'en' in automatic_captions:
                captions_available = True

            if captions_available:
                ydl.download([video_url])
                video_id = info_dict.get('id')
                caption_file = f"{video_id}.en.vtt"
                if os.path.exists(caption_file):
                    with open(caption_file, 'r', encoding='utf-8') as f:
                        captions_content = f.read()
                    os.remove(caption_file)
                    captions = convert_vtt_to_text(captions_content)
                    return captions
                else:
                    return None
            else:
                return None
    except Exception as e:
        logger.error(f"An error occurred while downloading captions: {e}")
        return None

def convert_vtt_to_text(vtt_content):
    try:
        vtt_file = StringIO(vtt_content)
        text = ''
        for caption in webvtt.read_buffer(vtt_file):
            text += caption.text + '\n'
        return text
    except Exception as e:
        logger.error(f"An error occurred while converting VTT to text: {e}")
        return ''

def transcribe_audio_with_whisper_api(audio_file_path):
    try:
        with open(audio_file_path, 'rb') as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )
        return transcript.text
    except Exception as e:
        logger.error(f"An error occurred during transcription: {e}")
        return None

def summarize_text(text):
    try:
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a tutorial maker. You create written tutorials that are easy to follow."},
                {
                    "role": "user",
                    "content": f"""
                    This is an audio transcript of a tutorial, DIY project, or recipe video. Please create a detailed, structured written tutorial in Markdown format based on this transcript.

                    **Structure**:
                    1. **Summary**: Start with a brief summary of the tutorial. Summarize the goal and the key steps involved.
                    2. **What You Will Need**: List all the materials, tools, or prerequisites the user will need. Use bullet points.
                    3. **Step-by-Step Instructions**: Break down the process into clear, ordered steps. Use headers (### Step 1: ..., ### Step 2: ...) and ordered lists where appropriate.
                        - Include sub-steps if necessary.
                        - Use bullet points for additional notes or tips.
                    4. **Additional Notes**: If applicable, include a section for tips, common issues, or additional resources.

                    **Markdown Syntax**:
                    - Use # for main sections (e.g., Summary, What You Will Need).
                    - Use ## for major steps or sections within the tutorial.
                    - Use ### for sub-steps.
                    - Use bullet points for lists and additional tips or notes.

                    **Tone and Language**:
                    - Write in a clear, concise, and friendly tone.
                    - Assume the reader has only basic knowledge of the subject, and provide explanations where necessary.

                    **Handling Non-Essential Content**:
                    - Exclude any non-instructive content such as jokes, personal anecdotes, or irrelevant tangents from the final tutorial.

                    **Error Handling**:
                    - If a part of the transcript is unclear or seems incorrect, indicate this in the tutorial with a note (e.g., [Note: This part of the audio was unclear]).

                    **Transcript**:
                    {text}
                    """
                }
            ]
        )

        summary = completion.choices[0].message.content.strip()
        return summary
    except Exception as e:
        logger.error(f"An error occurred during text summarization: {e}")
        return None

# Generate a confirmation token
def generate_confirmation_token(email):
    return serializer.dumps(email, salt='email-confirmation-salt')

# Confirm the token
def confirm_token(token, expiration=3600):
    try:
        email = serializer.loads(
            token,
            salt='email-confirmation-salt',
            max_age=expiration
        )
        return email
    except:
        return False

# Send confirmation email using Brevo
def send_confirmation_email(user_email, token):
    api_key = os.getenv("BREVO_API_KEY")
    
    configuration = sib_api_v3_sdk.Configuration()
    configuration.api_key['api-key'] = api_key
    api_instance = sib_api_v3_sdk.TransactionalEmailsApi(sib_api_v3_sdk.ApiClient(configuration))
    
    confirmation_url = url_for('confirm_email', token=token, _external=True)
    logger.debug(f"Confirmation URL: {confirmation_url}")

    send_smtp_email = sib_api_v3_sdk.SendSmtpEmail(
        to=[{"email": user_email}],
        template_id=1,  # The ID of the transactional email template you created
        params={"confirmation_link": confirmation_url}  # Pass the confirmation URL here
    )
    
    try:
        api_instance.send_transac_email(send_smtp_email)
        logger.info(f"Confirmation email sent to: {user_email}")
    except Exception as e:
        logger.error(f"Failed to send confirmation email: {e}")

# Send password reset email using Brevo
def send_password_reset_email(user_email, token):
    api_key = os.getenv("BREVO_API_KEY")
    
    configuration = sib_api_v3_sdk.Configuration()
    configuration.api_key['api-key'] = api_key
    api_instance = sib_api_v3_sdk.TransactionalEmailsApi(sib_api_v3_sdk.ApiClient(configuration))
    
    reset_url = url_for('reset_password', token=token, _external=True)
    logger.debug(f"Password Reset URL: {reset_url}")

    send_smtp_email = sib_api_v3_sdk.SendSmtpEmail(
        to=[{"email": user_email}],
        template_id=2,  # The ID of the password reset email template you created
        params={"reset_link": reset_url}  # Pass the reset URL here
    )
    
    try:
        api_instance.send_transac_email(send_smtp_email)
        logger.info(f"Password reset email sent to: {user_email}")
    except Exception as e:
        logger.error(f"Failed to send password reset email to {user_email}: {e}")

if __name__ == '__main__':
    app.run(debug=True)
